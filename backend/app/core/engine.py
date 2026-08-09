"""Pandoc 转换引擎实现"""

from __future__ import annotations

import asyncio
import hashlib
import os
import re
import shutil
import subprocess
import tempfile
import time
from collections.abc import Awaitable, Callable
from pathlib import Path
from typing import Any
from urllib.parse import unquote, urlparse
from uuid import UUID

import pypandoc
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from PIL import Image, ImageOps

from app.core.browser_check import edge_manager
from app.core.interfaces import ConversionEngine, ProgressCallback
from app.core.log import log
from app.core.mermaid_renderer import is_available as mermaid_renderer_available

# ── 嵌入的 Mermaid 渲染器 ─────────────────────────────
from app.core.mermaid_renderer import render_diagrams
from app.core.template_manager import TemplateManager
from app.models import ConversionResult, OutputFormat
from app.models.templates import ConversionOptions
from app.services.template_service import TemplateService
from app.utils.config import AppSettings
from app.utils.exceptions import ConversionError, PandocNotFoundError, UnsupportedFormatError

# ── 字号映射 ──────────────────────────────────────────────
SIZE_MAP: dict[str, float] = {
    "初号": 42,
    "小初": 36,
    "一号": 26,
    "小一": 24,
    "二号": 22,
    "小二": 18,
    "三号": 16,
    "小三": 15,
    "四号": 14,
    "小四": 12,
    "五号": 10.5,
    "小五": 9,
    "六号": 7.5,
    "小六": 6.5,
    "七号": 5.5,
    "八号": 5,
}

ALIGN_MAP: dict[str, WD_ALIGN_PARAGRAPH] = {
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
}

MERMAID_IMAGE_MARKER = "markflow-mermaid-diagram"
TITLE_PAGE_METADATA = "markflow-title-page"
PAGE_HEADER_METADATA = "markflow-page-header"
MIN_PDF_SIZE = 100
PDF_STABLE_CHECKS = 2


def _parse_size(raw: str | float) -> float | None:
    if isinstance(raw, (int, float)):
        return float(raw)
    s = str(raw).strip()
    if s in SIZE_MAP:
        return SIZE_MAP[s]
    num = "".join(c for c in s if c.isdigit() or c == ".")
    try:
        return float(num)
    except ValueError:
        return None


# ── 表格 XML 操作辅助 ─────────────────────────────────────
def _resolve_tc(cell_or_tc: object) -> Any:
    """统一入口：从 python-docx Cell 或 CT_Tc 元素获取 CT_Tc"""
    if hasattr(cell_or_tc, "_tc"):
        return cell_or_tc._tc  # type: ignore[attr-defined]
    return cell_or_tc  # 已是 CT_Tc


def _set_cell_border(cell_or_tc: object, edge: str, weight: float, color: str) -> None:
    """设置单元格单边边框"""
    tc = _resolve_tc(cell_or_tc)
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn("w:tcBorders"))
    if tcBorders is None:
        tcBorders = OxmlElement("w:tcBorders")
        tcPr.append(tcBorders)
    # 移除同边旧定义
    for old in tcBorders.findall(qn(f"w:{edge}")):
        tcBorders.remove(old)
    el = OxmlElement(f"w:{edge}")
    el.set(qn("w:val"), "single")
    el.set(qn("w:sz"), str(int(weight * 8)))  # 1pt = 8 eighths
    el.set(qn("w:space"), "0")
    el.set(qn("w:color"), color.lstrip("#"))
    tcBorders.append(el)


def _clear_cell_border(cell_or_tc: object, edge: str) -> None:
    """清除单元格单边边框"""
    tc = _resolve_tc(cell_or_tc)
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn("w:tcBorders"))
    if tcBorders is None:
        return
    for old in tcBorders.findall(qn(f"w:{edge}")):
        tcBorders.remove(old)


def _set_cell_margins(cell_or_tc: object, top: int, bottom: int, left: int, right: int) -> None:
    """设置单元格内边距（单位: dxa，1pt ≈ 20 dxa）"""
    tc = _resolve_tc(cell_or_tc)
    tcPr = tc.get_or_add_tcPr()
    for old in tcPr.findall(qn("w:tcMar")):
        tcPr.remove(old)
    tcMar = OxmlElement("w:tcMar")
    for direction, val in [("top", top), ("bottom", bottom), ("left", left), ("right", right)]:
        mar = OxmlElement(f"w:{direction}")
        mar.set(qn("w:w"), str(val * 20))  # dxa
        mar.set(qn("w:type"), "dxa")
        tcMar.append(mar)
    tcPr.append(tcMar)


def _set_cell_shading(cell_or_tc: object, fill_hex: str) -> None:
    """设置单元格底色"""
    tc = _resolve_tc(cell_or_tc)
    tcPr = tc.get_or_add_tcPr()
    for old in tcPr.findall(qn("w:shd")):
        tcPr.remove(old)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex.lstrip("#"))
    tcPr.append(shd)


def _apply_cell_text_format(
    cell_or_tc: object,
    font_name: str | None = None,
    font_size_pt: float | None = None,
    bold: bool | None = None,
    alignment: WD_ALIGN_PARAGRAPH | None = None,
) -> None:
    """对单元格内所有段落应用文本格式"""
    tc = _resolve_tc(cell_or_tc)
    for paragraph in tc.getchildren():  # type: ignore[attr-defined]
        p = paragraph  # w:p element
        if alignment is not None:
            pPr = p.find(qn("w:pPr"))
            if pPr is None:
                pPr = OxmlElement("w:pPr")
                p.insert(0, pPr)
            jc = pPr.find(qn("w:jc"))
            if jc is None:
                jc = OxmlElement("w:jc")
                pPr.append(jc)
            jc.set(
                qn("w:val"),
                {
                    WD_ALIGN_PARAGRAPH.CENTER: "center",
                    WD_ALIGN_PARAGRAPH.LEFT: "left",
                    WD_ALIGN_PARAGRAPH.RIGHT: "right",
                    WD_ALIGN_PARAGRAPH.JUSTIFY: "both",
                }.get(alignment, "left"),
            )

        for r in p.findall(qn("w:r")):
            rPr = r.find(qn("w:rPr"))
            if rPr is None:
                rPr = OxmlElement("w:rPr")
                r.insert(0, rPr)

            if font_name:
                rFonts = rPr.find(qn("w:rFonts"))
                if rFonts is None:
                    rFonts = OxmlElement("w:rFonts")
                    rPr.insert(0, rFonts)
                rFonts.set(qn("w:ascii"), font_name)
                rFonts.set(qn("w:eastAsia"), font_name)
                rFonts.set(qn("w:hAnsi"), font_name)
                for attr in (
                    qn("w:asciiTheme"),
                    qn("w:eastAsiaTheme"),
                    qn("w:hAnsiTheme"),
                    qn("w:cstheme"),
                ):
                    if attr in rFonts.attrib:
                        del rFonts.attrib[attr]

            if font_size_pt is not None:
                sz = rPr.find(qn("w:sz"))
                if sz is None:
                    sz = OxmlElement("w:sz")
                    rPr.append(sz)
                sz.set(qn("w:val"), str(int(font_size_pt * 2)))
                szCs = rPr.find(qn("w:szCs"))
                if szCs is not None:
                    rPr.remove(szCs)

            if bold is not None:
                existing_b = rPr.find(qn("w:b"))
                if bold:
                    if existing_b is None:
                        rPr.append(OxmlElement("w:b"))
                elif existing_b is not None:
                    rPr.remove(existing_b)


def _apply_three_line_border(table: object, tc: dict) -> None:
    """对整个表格应用三线表边框"""
    tbl = table._tbl  # type: ignore[attr-defined]
    rows = list(tbl.findall(qn("w:tr")))

    # 1) 表级属性：四边无边框 + 内无横竖线
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    for old in tblPr.findall(qn("w:tblBorders")):
        tblPr.remove(old)

    tblBorders = OxmlElement("w:tblBorders")
    _top_cfg = tc.get("border_top", {})
    _btm_cfg = tc.get("border_bottom", {})
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "none")
        b.set(qn("w:sz"), "0")
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), "auto")
        tblBorders.append(b)
    tblPr.append(tblBorders)

    if not rows:
        return

    # 2) 顶线 → 首行所有单元格的上边框
    top_w = _top_cfg.get("weight", 1.5)
    top_c = _top_cfg.get("color", "black")
    for cell_obj in rows[0].findall(qn("w:tc")):
        _set_cell_border(cell_obj, "top", top_w, top_c)

    # 3) 底线 → 末行所有单元格的下边框
    btm_w = _btm_cfg.get("weight", 1.5)
    btm_c = _btm_cfg.get("color", "black")
    for cell_obj in rows[-1].findall(qn("w:tc")):
        _set_cell_border(cell_obj, "bottom", btm_w, btm_c)

    # 4) 表头下框线 → 首行所有单元格的下边框
    hdr_cfg = tc.get("border_header_bottom", {})
    hdr_w = hdr_cfg.get("weight", 0.75)
    hdr_c = hdr_cfg.get("color", "black")
    for cell_obj in rows[0].findall(qn("w:tc")):
        _set_cell_border(cell_obj, "bottom", hdr_w, hdr_c)

    # 5) 清除其他所有单元格的上下左右边框
    no_h = tc.get("border_horizontal_internal", False)
    no_v = tc.get("border_vertical", False)
    for row_idx, row in enumerate(rows):
        for cell_obj in row.findall(qn("w:tc")):
            if no_h and row_idx > 0:
                _clear_cell_border(cell_obj, "top")
            if no_v:
                _clear_cell_border(cell_obj, "left")
                _clear_cell_border(cell_obj, "right")


def _apply_grid_border(table: object, tc: dict) -> None:
    """
    对表格应用全框线网格边框

    标准测试报告风格：所有单元格四边均为 0.5~0.75pt 黑色实线。
    """
    tbl = table._tbl  # type: ignore[attr-defined]
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    for old in tblPr.findall(qn("w:tblBorders")):
        tblPr.remove(old)

    # 各边框线宽取配置，默认 0.75pt → 6 个八分之一磅
    def _edge_weight(key: str, default: float = 0.75) -> int:
        cfg = tc.get(key, {})
        w = cfg.get("weight", default) if isinstance(cfg, dict) else default
        return max(1, int(w * 8))

    def _edge_color(key: str, default: str = "black") -> str:
        cfg = tc.get(key, {})
        c = cfg.get("color", default) if isinstance(cfg, dict) else default
        return str(c).lstrip("#")

    tblBorders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), str(_edge_weight(f"border_{edge}")))
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), _edge_color(f"border_{edge}"))
        tblBorders.append(b)
    tblPr.append(tblBorders)


class PandocEngine(ConversionEngine):
    """基于 Pandoc 的转换引擎（适配器模式）"""

    # OutputFormat → Pandoc -t 格式名
    FORMAT_MAP: dict[OutputFormat, str] = {
        OutputFormat.DOCX: "docx",
        OutputFormat.PDF: "pdf",
        OutputFormat.HTML: "html",
        OutputFormat.EPUB: "epub",
        OutputFormat.LATEX: "latex",
        OutputFormat.MARKDOWN: "gfm",
        OutputFormat.ODT: "odt",
        OutputFormat.RTF: "rtf",
    }

    def __init__(
        self,
        settings: AppSettings | None = None,
        template_manager: TemplateService | TemplateManager | None = None,
    ) -> None:
        self.settings = settings or AppSettings()
        self._template_mgr = template_manager or TemplateManager()
        # 不再在 __init__ 中强校验 Pandoc，改为在 convert() 时检查
        self._pandoc_available = self._check_pandoc()
        if not self._pandoc_available:
            log.warning("Pandoc 未安装，转换功能暂不可用。请在设置中安装 Pandoc 模块。")

    def _check_pandoc(self) -> bool:
        """检查 Pandoc 是否可用（非强制）"""
        try:
            path = pypandoc.get_pandoc_path()
            log.info(f"Pandoc 路径: {path}")
            return True
        except OSError as e:
            log.debug(f"Pandoc 检查: {e}")
            return False
        except Exception:
            return False

    async def convert(  # noqa: PLR0913
        self,
        input_path: Path,
        output_format: OutputFormat,
        extra_args: list[str] | None = None,
        template_slug: str | None = None,
        *,
        convert_images: bool = True,
        convert_mermaid: bool = True,
        options: dict | None = None,
        on_progress: ProgressCallback | None = None,
    ) -> ConversionResult:
        """执行 Pandoc 转换"""
        del options
        # 运行时检查 Pandoc 是否可用
        if not self._pandoc_available:
            self._pandoc_available = self._check_pandoc()
        if not self._pandoc_available:
            raise PandocNotFoundError(
                "Pandoc 未安装，请在设置中安装 Pandoc 模块后再试。",
                detail={"hint": "打开设置 → 模块 → Pandoc → 安装"},
            )

        pandoc_target = self.FORMAT_MAP.get(output_format)
        if pandoc_target is None:
            raise UnsupportedFormatError(f"不支持的格式: {output_format.value}")

        if on_progress:
            await on_progress(0.05, "准备转换...")

        # 预处理 Markdown：标准化列表结构（确保列表前有空行）
        self._normalize_markdown_structure(input_path)
        # 预处理 Markdown：标准化数学公式定界符
        self._normalize_math_in_file(input_path)
        # 预处理 Markdown：渲染 Mermaid 图表
        created_dirs = (
            await self._preprocess_mermaid(
                input_path,
                mark_for_docx=output_format == OutputFormat.DOCX,
            )
            if convert_mermaid
            else []
        )
        created_dirs.extend(self._preprocess_images(input_path, convert_images))

        output_path = input_path.with_suffix(f".{output_format.value}")
        args = list(extra_args or [])

        # 组装模版参数
        # API 层已经按完整 ConversionOptions 组装过参数时，不再重复追加 reference/filter。
        # 直接调用引擎且未提供模板参数时，仍保留原有的自动解析行为。
        if template_slug and "--reference-doc" not in args:
            template_args = self._template_mgr.build_extra_args(
                ConversionOptions(template_slug=template_slug),
            )
            args = template_args + args
            log.info(f"使用模版: {template_slug}, 参数: {template_args}")

        title_page = self._get_metadata_value(args, TITLE_PAGE_METADATA) == "true"
        page_header = self._get_metadata_value(args, PAGE_HEADER_METADATA)
        pandoc_args = self._remove_metadata_keys(
            args,
            {TITLE_PAGE_METADATA, PAGE_HEADER_METADATA},
        )

        start = time.monotonic()

        if on_progress:
            await on_progress(0.1, "正在转换...")

        try:
            loop = asyncio.get_running_loop()

            def _convert_with_pandoc(
                target: str,
                destination: Path,
                conversion_args: list[str],
            ) -> str:
                return pypandoc.convert_file(
                    source_file=str(input_path),
                    to=target,
                    format="markdown",
                    outputfile=str(destination),
                    extra_args=conversion_args,
                )

            if output_format == OutputFormat.PDF:
                html_path = input_path.with_name(f"{input_path.stem}.markflow-pdf.html")
                css_path = input_path.with_name(f"{input_path.stem}.markflow-pdf.css")
                css_path.write_text(
                    self._build_pdf_css(template_slug, title_page, page_header),
                    encoding="utf-8",
                )
                pdf_args = self._remove_pandoc_options(
                    pandoc_args,
                    {"--reference-doc", "--pdf-engine"},
                )
                pdf_args.extend(
                    [
                        "--standalone",
                        "--embed-resources",
                        "--mathml",
                        "--css",
                        str(css_path.resolve()),
                    ]
                )
                try:
                    await self._run_with_timeout(
                        loop.run_in_executor(
                            None,
                            _convert_with_pandoc,
                            "html5",
                            html_path,
                            pdf_args,
                        )
                    )
                    if on_progress:
                        await on_progress(0.75, "正在生成 PDF...")
                    await self._render_html_to_pdf(html_path, output_path)
                finally:
                    html_path.unlink(missing_ok=True)
                    css_path.unlink(missing_ok=True)
            else:
                await self._run_with_timeout(
                    loop.run_in_executor(
                        None,
                        _convert_with_pandoc,
                        pandoc_target,
                        output_path,
                        pandoc_args,
                    )
                )

            # 后处理：应用表格样式（仅 docx）
            # Pandoc 在 DOCX 中只写入空的 TOC 域。填充可见缓存，并让 Word
            # 打开文档时自动更新目录和页码。
            if output_format == OutputFormat.DOCX and "--toc" in pandoc_args:
                toc_depth = self._get_toc_depth(args)
                if on_progress:
                    await on_progress(0.88, "生成目录...")
                await loop.run_in_executor(
                    None,
                    self._populate_docx_toc_cache,
                    output_path,
                    toc_depth,
                )

            if output_format == OutputFormat.DOCX and (title_page or page_header):
                if on_progress:
                    await on_progress(0.89, "设置标题页和页眉...")
                header_config = (
                    self._template_mgr.get_header_config(template_slug)
                    if template_slug
                    else None
                )
                await loop.run_in_executor(
                    None,
                    self._apply_docx_page_options,
                    output_path,
                    title_page,
                    page_header,
                    header_config,
                )

            if output_format == OutputFormat.DOCX and template_slug:
                if on_progress:
                    await on_progress(0.9, "应用表格样式...")
                await loop.run_in_executor(
                    None, self._apply_table_styles, output_path, template_slug
                )

            if output_format == OutputFormat.DOCX:
                if on_progress:
                    await on_progress(0.95, "调整 Mermaid 图表版式...")
                await loop.run_in_executor(
                    None,
                    self._format_docx_mermaid_images,
                    output_path,
                )

            if on_progress:
                await on_progress(1.0, "转换完成")

            duration = int((time.monotonic() - start) * 1000)
            file_size = output_path.stat().st_size

            log.info(
                f"转换完成: {input_path.name} → {output_format.value}, "
                f"耗时 {duration}ms, 大小 {file_size} bytes"
            )

            return ConversionResult(
                task_id=UUID("00000000-0000-0000-0000-000000000000"),
                output_path=output_path,
                output_format=output_format,
                duration_ms=duration,
                file_size=file_size,
            )

        except TimeoutError:
            raise ConversionError(
                f"转换超时（{self.settings.pandoc_timeout}s）",
                detail={"input": str(input_path), "format": output_format.value},
            ) from None
        except Exception as e:
            log.error(f"Pandoc 转换失败: {input_path.name} → {output_format.value}: {e}")
            raise ConversionError(
                "Pandoc 转换失败",
                detail={
                    "input": str(input_path),
                    "format": output_format.value,
                    "error": str(e),
                },
            ) from e
        finally:
            # 清理 Mermaid 和图片预处理产生的临时目录
            for d in created_dirs:
                self.__class__._cleanup_mermaid_dir(d)

    async def validate_format(self, output_format: OutputFormat) -> bool:
        """校验格式是否支持"""
        return output_format in self.FORMAT_MAP

    async def _run_with_timeout(self, awaitable: Awaitable[Any]) -> Any:
        if self.settings.pandoc_timeout > 0:
            return await asyncio.wait_for(
                awaitable,
                timeout=self.settings.pandoc_timeout,
            )
        return await awaitable

    @staticmethod
    def _remove_pandoc_options(args: list[str], options: set[str]) -> list[str]:
        """移除不适用于中间 HTML 的 Pandoc 参数及其值。"""
        filtered: list[str] = []
        index = 0
        while index < len(args):
            arg = args[index]
            option = arg.partition("=")[0]
            if option in options:
                index += 1 if "=" in arg else 2
                continue
            filtered.append(arg)
            index += 1
        return filtered

    async def _render_html_to_pdf(self, html_path: Path, output_path: Path) -> None:
        """使用系统 Edge 将自包含 HTML 打印为 PDF。"""
        edge = edge_manager.executable_path()
        if not edge:
            raise ConversionError(
                "未找到 Microsoft Edge，无法导出 PDF",
                detail={"hint": "请安装或修复 Microsoft Edge 后重试"},
            )

        # Edge 偶尔会把打印任务交给子进程后提前退出。若目标文件残留，
        # 后续的存在性检查还会把旧文件误判成这次转换的结果。
        output_path.unlink(missing_ok=True)
        pdf_ready = False
        with tempfile.TemporaryDirectory(
            prefix="markflow-pdf-edge-",
            dir=html_path.parent,
        ) as profile_dir:
            profile_path = Path(profile_dir).resolve()
            creationflags = 0
            if os.name == "nt":
                creationflags = subprocess.CREATE_NO_WINDOW  # type: ignore[attr-defined]
            process = await asyncio.create_subprocess_exec(
                edge,
                "--headless=new",
                "--disable-gpu",
                "--no-sandbox",
                "--disable-breakpad",
                "--disable-crash-reporter",
                "--no-pdf-header-footer",
                "--print-to-pdf-no-header",
                f"--user-data-dir={profile_path}",
                f"--print-to-pdf={output_path.resolve()}",
                html_path.resolve().as_uri(),
                stdout=asyncio.subprocess.DEVNULL,
                stderr=asyncio.subprocess.PIPE,
                creationflags=creationflags,
            )
            try:
                _, stderr = await asyncio.wait_for(
                    process.communicate(),
                    timeout=self.settings.pandoc_timeout or 300,
                )
            except TimeoutError:
                process.kill()
                await process.communicate()
                raise
            if process.returncode == 0:
                # 打包后的 sidecar 中较容易遇到 Edge 进程已退出、PDF 仍在
                # 后台落盘的情况。临时配置和 HTML 必须保留到文件写完。
                pdf_ready = await self._wait_for_pdf_output(output_path)

        if process.returncode != 0 or not pdf_ready:
            detail = stderr.decode("utf-8", errors="replace").strip()[:500]
            log.warning(
                f"Edge 生成 PDF 失败 "
                f"(exit={process.returncode}, ready={pdf_ready}, "
                f"exists={output_path.exists()}): {detail}"
            )
            raise ConversionError(
                "Edge 生成 PDF 失败",
                detail={"exit_code": process.returncode, "error": detail},
            )

    @staticmethod
    async def _wait_for_pdf_output(
        output_path: Path,
        timeout_seconds: float = 5.0,
    ) -> bool:
        """等待 Edge 的异步打印子进程完成写入。"""
        loop = asyncio.get_running_loop()
        deadline = loop.time() + timeout_seconds
        previous_size = -1
        stable_checks = 0

        while loop.time() < deadline:
            try:
                size = output_path.stat().st_size
                with output_path.open("rb") as stream:
                    is_pdf = stream.read(5) == b"%PDF-"
            except OSError:
                size = 0
                is_pdf = False

            if is_pdf and size >= MIN_PDF_SIZE:
                if size == previous_size:
                    stable_checks += 1
                    if stable_checks >= PDF_STABLE_CHECKS:
                        return True
                else:
                    previous_size = size
                    stable_checks = 0
            else:
                previous_size = size
                stable_checks = 0

            await asyncio.sleep(0.1)

        return False

    def _build_pdf_css(
        self,
        template_slug: str | None,
        title_page: bool,
        page_header: str,
    ) -> str:
        """将模板的常用 Word 样式映射为可打印 CSS。"""
        styles = (
            self._template_mgr.get_styles_config(template_slug)
            if template_slug
            else {}
        )
        body = styles.get("body", {})
        code = styles.get("code", {})
        table = styles.get("table", {})

        body_font = body.get("font", "宋体")
        body_size = _parse_size(body.get("size", "小四")) or 12
        line_spacing = body.get("line_spacing", 1.5)
        first_indent = "2em" if body.get("first_line_indent") else "0"
        code_font = code.get("font", "Consolas")
        code_size = _parse_size(code.get("size", "五号")) or 10.5
        code_background = code.get("background", "#f5f5f5")
        header_css = ""
        if page_header:
            safe_header = page_header.replace("\\", "\\\\").replace('"', '\\"')
            header_css = (
                'body::before { content: "'
                + safe_header
                + '"; position: fixed; top: -13mm; left: 0; right: 0; '
                "text-align: center; font-size: 10.5pt; color: #555; "
                "border-bottom: .75pt solid #777; padding-bottom: 2mm; }"
            )

        heading_rules: list[str] = []
        for level in range(1, 7):
            config = styles.get(f"heading{level}", {})
            fallback_size = max(12, 22 - level * 2)
            size = _parse_size(config.get("size", fallback_size)) or fallback_size
            font = config.get("font", body_font)
            color = config.get("color", "#111")
            alignment = config.get("alignment", "left")
            weight = "700" if config.get("bold", True) else "400"
            heading_rules.append(
                f"h{level} {{ font-family: {self._css_font(font)}; font-size: {size}pt; "
                f"font-weight: {weight}; color: {color}; text-align: {alignment}; "
                "line-height: 1.35; break-after: avoid; }"
            )

        title_break = (
            ".title-block-header { min-height: 230mm; display: flex; "
            "flex-direction: column; justify-content: center; text-align: center; "
            "break-after: page; }"
            if title_page
            else ""
        )
        stripe = (
            f"tbody tr:nth-child(even) {{ background: {table.get('stripe_color', '#f7f7f7')}; }}"
            if table.get("stripe_rows")
            else ""
        )
        return "\n".join(
            [
                "@page { size: A4; margin: 22mm 20mm 20mm; }",
                "html { print-color-adjust: exact; -webkit-print-color-adjust: exact; }",
                (
                    f"body {{ font-family: {self._css_font(body_font)}; "
                    f"font-size: {body_size}pt; line-height: {line_spacing}; "
                    "color: #222; overflow-wrap: anywhere; }}"
                ),
                "p { margin: .45em 0; }",
                f"p:not(.author):not(.date) {{ text-indent: {first_indent}; }}",
                *heading_rules,
                "img, svg { display: block; max-width: 100%; height: auto; margin: 1em auto; }",
                "figure { margin: 1em 0; break-inside: avoid; }",
                "table { width: 100%; border-collapse: collapse; margin: 1em 0; "
                "font-size: .9em; break-inside: auto; }",
                "thead { display: table-header-group; } tr { break-inside: avoid; }",
                "th, td { padding: 5pt 7pt; border-bottom: .5pt solid #bbb; }",
                f"th {{ background: {table.get('header_background', '#f2f2f2')}; "
                "font-weight: 700; }",
                stripe,
                (
                    f"pre, code {{ font-family: {self._css_font(code_font)}; "
                    f"font-size: {code_size}pt; }}"
                ),
                f"pre {{ background: {code_background}; padding: 10pt; border-radius: 4pt; "
                "white-space: pre-wrap; break-inside: avoid; }",
                "blockquote { margin: 1em 0; padding: .2em 1em; color: #555; "
                "border-left: 3pt solid #bbb; }",
                "a { color: inherit; text-decoration: none; }",
                ".math.display { display: block; text-align: center; margin: 1em 0; }",
                title_break,
                header_css,
            ]
        )

    @staticmethod
    def _css_font(font: object) -> str:
        name = str(font).replace("\\", "\\\\").replace('"', '\\"')
        return f'"{name}", "Microsoft YaHei", "Noto Sans CJK SC", sans-serif'

    @staticmethod
    def _get_toc_depth(args: list[str]) -> int:
        for index, arg in enumerate(args):
            if arg == "--toc-depth" and index + 1 < len(args):
                try:
                    return max(1, min(6, int(args[index + 1])))
                except ValueError:
                    return 3
            if arg.startswith("--toc-depth="):
                try:
                    return max(1, min(6, int(arg.partition("=")[2])))
                except ValueError:
                    return 3
        return 3

    @staticmethod
    def _get_metadata_value(args: list[str], key: str) -> str:
        """读取 ``--metadata key=value`` 形式的 Pandoc 参数。"""
        for index, arg in enumerate(args):
            value = ""
            if arg == "--metadata" and index + 1 < len(args):
                value = args[index + 1]
            elif arg.startswith("--metadata="):
                value = arg.partition("=")[2]
            if value.partition("=")[0] == key:
                return value.partition("=")[2]
        return ""

    @staticmethod
    def _remove_metadata_keys(args: list[str], keys: set[str]) -> list[str]:
        """移除仅供 MarkFlow 后处理使用的内部元数据参数。"""
        filtered: list[str] = []
        index = 0
        while index < len(args):
            arg = args[index]
            if arg == "--metadata" and index + 1 < len(args):
                value = args[index + 1]
                if value.partition("=")[0] in keys:
                    index += 2
                    continue
            elif arg.startswith("--metadata="):
                value = arg.partition("=")[2]
                if value.partition("=")[0] in keys:
                    index += 1
                    continue
            filtered.append(arg)
            index += 1
        return filtered

    @staticmethod
    def _apply_docx_page_options(
        docx_path: Path,
        title_page: bool,
        page_header: str,
        header_config: dict | None = None,
    ) -> None:
        """为 DOCX 设置独立标题页和顶部页眉。"""
        doc = Document(docx_path)

        if title_page:
            title_style_names = {"Title", "Subtitle", "Author", "Date"}
            last_title_paragraph = None
            found_title = False
            for paragraph in doc.paragraphs:
                style_name = paragraph.style.name if paragraph.style is not None else ""
                if style_name in title_style_names:
                    found_title = True
                    last_title_paragraph = paragraph
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    continue
                if found_title and not paragraph.text.strip():
                    last_title_paragraph = paragraph
                    continue
                if found_title:
                    break

            if last_title_paragraph is not None:
                next_element = last_title_paragraph._p.getnext()
                next_paragraph = None
                while next_element is not None and next_paragraph is None:
                    if next_element.tag == qn("w:p"):
                        next_paragraph = next_element
                    else:
                        next_paragraph = next_element.find(".//" + qn("w:p"))
                    next_element = next_element.getnext()
                if next_paragraph is not None:
                    PandocEngine._set_page_break_before(next_paragraph)
                else:
                    log.warning(
                        f"DOCX 中标题后没有可分页的内容: {docx_path.name}"
                    )
            else:
                log.warning(f"DOCX 中未找到标题段落，无法生成独立标题页: {docx_path.name}")

        normalized_header = page_header.strip()
        config = header_config if isinstance(header_config, dict) else {}
        font_name = str(config.get("font", "宋体"))
        font_size = _parse_size(config.get("size", "五号")) or 10.5
        alignment = ALIGN_MAP.get(
            str(config.get("alignment", "center")),
            WD_ALIGN_PARAGRAPH.CENTER,
        )
        raw_border_config = config.get("border_bottom", {})
        border_config = raw_border_config if isinstance(raw_border_config, dict) else {}
        border_weight = float(border_config.get("weight", 0.75))
        border_color = str(border_config.get("color", "000000")).lstrip("#")
        if border_color.lower() == "black":
            border_color = "000000"

        for section_index, section in enumerate(doc.sections):
            if title_page and section_index == 0:
                # 标题页也使用正常页眉，避免 Word 把首页页眉隐藏。
                section.different_first_page_header_footer = False
            if not normalized_header:
                continue
            header = section.header
            paragraph = header.paragraphs[0]
            paragraph.clear()
            paragraph.paragraph_format.left_indent = Pt(0)
            paragraph.paragraph_format.right_indent = Pt(0)
            paragraph.alignment = alignment

            run = paragraph.add_run(normalized_header)
            run.font.name = font_name
            run.font.size = Pt(font_size)
            run_properties = run._r.get_or_add_rPr()
            run_fonts = run_properties.find(qn("w:rFonts"))
            if run_fonts is None:
                run_fonts = OxmlElement("w:rFonts")
                run_properties.insert(0, run_fonts)
            for attribute in ("ascii", "eastAsia", "hAnsi"):
                run_fonts.set(qn(f"w:{attribute}"), font_name)

            paragraph_properties = paragraph._p.get_or_add_pPr()
            borders = paragraph_properties.find(qn("w:pBdr"))
            if borders is None:
                borders = OxmlElement("w:pBdr")
                paragraph_properties.append(borders)
            for old_bottom in borders.findall(qn("w:bottom")):
                borders.remove(old_bottom)
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"), "single")
            bottom.set(qn("w:sz"), str(max(2, round(border_weight * 8))))
            bottom.set(qn("w:space"), "1")
            bottom.set(qn("w:color"), border_color)
            borders.append(bottom)

        doc.save(docx_path)

    @staticmethod
    def _set_page_break_before(paragraph: object) -> None:
        """使用段前分页，避免 Word 显示可见的“分页符”标记。"""
        paragraph_properties = paragraph.find(qn("w:pPr"))
        if paragraph_properties is None:
            paragraph_properties = OxmlElement("w:pPr")
            paragraph.insert(0, paragraph_properties)
        page_break = paragraph_properties.find(qn("w:pageBreakBefore"))
        if page_break is None:
            page_break = OxmlElement("w:pageBreakBefore")
            paragraph_properties.append(page_break)
        page_break.set(qn("w:val"), "1")

    @staticmethod
    def _populate_docx_toc_cache(docx_path: Path, toc_depth: int) -> None:
        """为 Pandoc 生成的空 TOC 域写入可见标题缓存。"""
        doc = Document(docx_path)
        body = doc.element.body

        toc_content = None
        for sdt in body.findall(qn("w:sdt")):
            gallery = sdt.find(".//" + qn("w:docPartGallery"))
            if gallery is not None and gallery.get(qn("w:val")) == "Table of Contents":
                toc_content = sdt.find(qn("w:sdtContent"))
                break

        if toc_content is None:
            log.warning(f"DOCX 中未找到目录域: {docx_path.name}")
            return

        headings: list[tuple[int, str, str | None]] = []
        pending_bookmark: str | None = None
        for child in body:
            if child.tag == qn("w:bookmarkStart"):
                pending_bookmark = child.get(qn("w:name"))
                continue
            if child.tag != qn("w:p"):
                continue

            p_style = child.find("./" + qn("w:pPr") + "/" + qn("w:pStyle"))
            style_value = p_style.get(qn("w:val"), "") if p_style is not None else ""
            match = re.fullmatch(r"Heading([1-6])", style_value)
            if match:
                level = int(match.group(1))
                title = "".join(node.text or "" for node in child.iter(qn("w:t"))).strip()
                if title and level <= toc_depth:
                    headings.append((level, title, pending_bookmark))
            pending_bookmark = None

        field_paragraph = None
        for paragraph in toc_content.findall(qn("w:p")):
            instruction = "".join(
                node.text or "" for node in paragraph.iter(qn("w:instrText"))
            )
            if "TOC " in instruction:
                field_paragraph = paragraph
                break

        if field_paragraph is None:
            log.warning(f"DOCX 中未找到 TOC 指令: {docx_path.name}")
            return

        # 清除可能存在的旧缓存，并把域结束标记移动到缓存条目之后。
        field_index = toc_content.index(field_paragraph)
        for child in list(toc_content)[field_index + 1 :]:
            toc_content.remove(child)
        for field_char in list(field_paragraph.iter(qn("w:fldChar"))):
            if field_char.get(qn("w:fldCharType")) == "end":
                field_char.getparent().remove(field_char)

        insert_at = field_index + 1
        for level, title, bookmark in headings:
            paragraph = OxmlElement("w:p")
            paragraph_properties = OxmlElement("w:pPr")
            indent = OxmlElement("w:ind")
            indent.set(qn("w:left"), str((level - 1) * 420))
            paragraph_properties.append(indent)
            spacing = OxmlElement("w:spacing")
            spacing.set(qn("w:before"), "0")
            spacing.set(qn("w:after"), "0")
            paragraph_properties.append(spacing)
            paragraph.append(paragraph_properties)

            run = OxmlElement("w:r")
            text = OxmlElement("w:t")
            text.set(qn("xml:space"), "preserve")
            text.text = title
            run.append(text)

            if bookmark:
                hyperlink = OxmlElement("w:hyperlink")
                hyperlink.set(qn("w:anchor"), bookmark)
                hyperlink.set(qn("w:history"), "1")
                hyperlink.append(run)
                paragraph.append(hyperlink)
            else:
                paragraph.append(run)

            toc_content.insert(insert_at, paragraph)
            insert_at += 1

        end_paragraph = OxmlElement("w:p")
        end_run = OxmlElement("w:r")
        end_field = OxmlElement("w:fldChar")
        end_field.set(qn("w:fldCharType"), "end")
        end_run.append(end_field)
        end_paragraph.append(end_run)
        toc_content.insert(insert_at, end_paragraph)

        # 目录必须独占页面：对 TOC 后的首个正文段落设置“段前分页”。
        # 这不会像手动分页符那样在 Word 中显示“——分页符——”标记。
        toc_container = toc_content.getparent()
        if toc_container is not None:
            toc_index = body.index(toc_container)
            for following in list(body)[toc_index + 1 :]:
                if following.tag != qn("w:p"):
                    continue
                PandocEngine._set_page_break_before(following)
                break

        settings = doc.settings.element
        update_fields = settings.find(qn("w:updateFields"))
        if update_fields is None:
            update_fields = OxmlElement("w:updateFields")
            settings.append(update_fields)
        update_fields.set(qn("w:val"), "true")

        doc.save(docx_path)
        log.info(f"已生成可见目录缓存: {len(headings)} 项, 深度 {toc_depth}")

    # ── Markdown 结构预处理 ──────────────────────────────

    @staticmethod
    def _format_docx_mermaid_images(docx_path: Path) -> None:
        """Center Mermaid figures and constrain them to the current Word section."""
        try:
            doc = Document(docx_path)
        except Exception as exc:
            log.warning(f"无法打开 DOCX 调整 Mermaid 图表版式: {exc}")
            return
        sections = list(doc.sections)
        if not sections:
            return

        body = doc.element.body
        paragraph_sections: dict[object, int] = {}
        section_index = 0
        for child in body.iterchildren():
            if child.tag != qn("w:p"):
                continue
            paragraph_sections[child] = min(section_index, len(sections) - 1)
            if child.find("./" + qn("w:pPr") + "/" + qn("w:sectPr")) is not None:
                section_index += 1

        formatted = 0
        for shape in doc.inline_shapes:
            inline = shape._inline
            doc_properties = inline.find(qn("wp:docPr"))
            if doc_properties is None:
                continue
            metadata = " ".join(
                doc_properties.get(key, "")
                for key in ("name", "title", "descr")
            ).lower()
            if MERMAID_IMAGE_MARKER not in metadata:
                continue

            paragraph_element = inline
            while (
                paragraph_element is not None
                and paragraph_element.tag != qn("w:p")
            ):
                paragraph_element = paragraph_element.getparent()
            if paragraph_element is None:
                continue

            target_section = sections[
                paragraph_sections.get(paragraph_element, len(sections) - 1)
            ]
            page_width = int(target_section.page_width or Inches(8.5))
            left_margin = int(target_section.left_margin or Inches(1))
            right_margin = int(target_section.right_margin or Inches(1))
            available_width = page_width - left_margin - right_margin
            if available_width <= 0:
                continue

            current_width = int(shape.width)
            current_height = int(shape.height)
            if current_width > available_width:
                ratio = available_width / current_width
                shape.width = available_width
                shape.height = round(current_height * ratio)

            paragraph_properties = paragraph_element.get_or_add_pPr()
            for old_indent in paragraph_properties.findall(qn("w:ind")):
                paragraph_properties.remove(old_indent)
            justification = paragraph_properties.find(qn("w:jc"))
            if justification is None:
                justification = OxmlElement("w:jc")
                paragraph_properties.append(justification)
            justification.set(qn("w:val"), "center")
            formatted += 1

        if formatted:
            doc.save(docx_path)
            log.info(f"已居中并适配 {formatted} 张 Mermaid Word 图表")

    @staticmethod
    def _normalize_markdown_structure(path: Path) -> None:
        """
        标准化 Markdown 结构：确保列表前有空行

        Pandoc 的 markdown 解析要求列表前有空行，否则列表标记 `-`、`*` 等
        会被当作普通文本中的连字符/星号处理，导致结构被拍平。

        预处理扫描所有行，如果某行是列表项且前一行非空且非列表项，
        则在其前插入空行，保证 Pandoc 能正确识别列表结构。
        """
        try:
            text = path.read_text(encoding="utf-8")
        except OSError:
            return

        # 统一换行符
        text = text.replace("\r\n", "\n").replace("\r", "\n")

        lines = text.split("\n")
        list_pattern = re.compile(r"^\s*([-*+]|\d+\.)\s")

        result: list[str] = []
        for i, line in enumerate(lines):
            stripped = line.strip()
            if stripped:  # 非空行
                is_list = bool(list_pattern.match(line))
                if is_list and i > 0:
                    prev_stripped = lines[i - 1].strip()
                    prev_is_list = prev_stripped and bool(list_pattern.match(lines[i - 1]))
                    if prev_stripped and not prev_is_list:
                        result.append("")  # 插入空行
            result.append(line)

        new_text = "\n".join(result)

        if new_text != text:
            try:
                path.write_text(new_text, encoding="utf-8")
                log.info(f"已标准化列表结构: {path.name}")
            except OSError:
                pass

    # ── 数学公式预处理 ───────────────────────────────────

    @staticmethod
    def _normalize_math_in_file(path: Path) -> None:
        r"""
        将 Markdown 中 `[ ... ]` 和 `\[...\]` 显示公式转为 Pandoc 标准语法 `$$ ... $$`

        Pandoc 只识别 $$...$$ 和 \\[...\\] 作为显示公式，
        许多作者习惯用 [ ... ] 或 \\[...\\]，需要提前归一化。

        支持四种格式：
        1. 单行：\\[ I_1 = K_3 I_N \\]
        2. 多行：\\[ \\n content \\n \\]
        3. 单行：[ I_1 = K_3 I_N ]
        4. 多行：[ \\n content \\n ]
        """
        try:
            text = path.read_text(encoding="utf-8")
        except OSError:
            return

        # 统一换行符（Windows \r\n → \n），否则 $ 锚点不匹配
        text = text.replace("\r\n", "\n").replace("\r", "\n")

        new_text = text

        # ── 多行显示公式（反斜杠格式）──
        # 匹配：独占一行的 \[ ，到独占一行的 \] ，中间为公式内容
        new_text = re.sub(
            r"^[ \t]*\\\[[ \t]*$\n(.*?)^[ \t]*\\\][ \t]*$",
            lambda m: "$$\n" + m.group(1) + "$$",
            new_text,
            flags=re.MULTILINE | re.DOTALL,
        )

        # ── 单行显示公式（反斜杠格式）──
        # 匹配：独占一行的 \[ math content \]
        new_text = re.sub(
            r"^[ \t]*\\\[[ \t]*(.+?)[ \t]*\\\][ \t]*$",
            r"$$ \1 $$",
            new_text,
            flags=re.MULTILINE,
        )

        # ── 多行显示公式（无反斜杠格式）──
        # 匹配：独占一行的 [ ，到独占一行的 ] ，中间为公式内容
        new_text = re.sub(
            r"^[ \t]*\[[ \t]*$\n(.*?)^[ \t]*\][ \t]*$",
            lambda m: "$$\n" + m.group(1) + "$$",
            new_text,
            flags=re.MULTILINE | re.DOTALL,
        )

        # ── 单行显示公式（无反斜杠格式）──
        # 匹配：[ math content ] 整个在一行
        new_text = re.sub(
            r"^[ \t]*\[[ \t]*(.+?(?:[_\\{}]|[a-z]+\^|sum|int|lim|prod|frac|sqrt|sin|cos|log).+?)[ \t]*\][ \t]*$",  # noqa: E501
            r"$$ \1 $$",
            new_text,
            flags=re.MULTILINE,
        )

        if new_text != text:
            try:
                path.write_text(new_text, encoding="utf-8")
                log.info(f"已标准化数学公式: {path.name}")
            except OSError:
                pass

    # ── 图片预处理 ────────────────────────────────────────

    _FENCED_CODE_RE = re.compile(
        r"(^[ \t]*```.*?^[ \t]*```[ \t]*$|^[ \t]*~~~.*?^[ \t]*~~~[ \t]*$)",
        re.MULTILINE | re.DOTALL,
    )
    _MARKDOWN_IMAGE_RE = re.compile(
        r"!\[(?P<alt>[^\]]*)\]\(\s*(?P<target><[^>]+>|[^\s)]+)"
        r"(?P<title>\s+(?:\"[^\"]*\"|'[^']*'|\([^)]*\)))?\s*\)",
    )
    _REFERENCE_IMAGE_RE = re.compile(r"!\[([^\]]*)\]\[[^\]]*\]")
    _HTML_IMAGE_RE = re.compile(r"<img\b[^>]*>", re.IGNORECASE)
    _ALT_ATTR_RE = re.compile(
        r"\balt\s*=\s*([\"'])(.*?)\1",
        re.IGNORECASE | re.DOTALL,
    )

    @classmethod
    def _map_outside_fenced_code(
        cls,
        text: str,
        transform: Callable[[str], str],
    ) -> str:
        parts = cls._FENCED_CODE_RE.split(text)
        return "".join(
            part if index % 2 else transform(part)
            for index, part in enumerate(parts)
        )

    @classmethod
    def _strip_images_from_segment(cls, segment: str) -> str:
        segment = cls._MARKDOWN_IMAGE_RE.sub(
            lambda match: match.group("alt"),
            segment,
        )
        segment = cls._REFERENCE_IMAGE_RE.sub(lambda match: match.group(1), segment)

        def replace_html(match: re.Match[str]) -> str:
            alt = cls._ALT_ATTR_RE.search(match.group(0))
            return alt.group(2) if alt else ""

        return cls._HTML_IMAGE_RE.sub(replace_html, segment)

    @staticmethod
    def _resolve_local_image(target: str, input_path: Path) -> Path | None:
        raw_target = target[1:-1] if target.startswith("<") else target
        parsed = urlparse(raw_target)
        if parsed.scheme.lower() in {"http", "https", "data"}:
            return None
        if parsed.scheme and parsed.scheme.lower() != "file":
            is_windows_path = (
                len(parsed.scheme) == 1 and raw_target[1:3] in {":\\", ":/"}
            )
            if not is_windows_path:
                return None

        if parsed.scheme.lower() == "file":
            value = unquote(parsed.path)
            if re.match(r"^/[A-Za-z]:/", value):
                value = value[1:]
        else:
            value = unquote(raw_target)

        path = Path(value)
        if not path.is_absolute():
            path = input_path.parent / path
        return path.resolve()

    @classmethod
    def _preprocess_images(
        cls,
        input_path: Path,
        convert_images: bool,
    ) -> list[Path]:
        try:
            text = input_path.read_text(encoding="utf-8")
        except OSError:
            return []

        if not convert_images:
            new_text = cls._map_outside_fenced_code(
                text,
                cls._strip_images_from_segment,
            )
            if new_text != text:
                input_path.write_text(new_text, encoding="utf-8")
                log.info(f"已从转换内容中移除图片: {input_path.name}")
            return []

        convertible = {".bmp", ".gif", ".tif", ".tiff", ".webp"}
        stem_hash = hashlib.sha256(str(input_path).encode("utf-8")).hexdigest()[:8]
        tmp_dir = input_path.parent / f"_images_{stem_hash}"
        converted: dict[Path, Path] = {}

        def replace_image(match: re.Match[str]) -> str:
            source = cls._resolve_local_image(match.group("target"), input_path)
            if source is None or source.suffix.lower() not in convertible:
                return match.group(0)
            if not source.is_file():
                log.warning(f"图片文件不存在，跳过转换: {source}")
                return match.group(0)

            output = converted.get(source)
            if output is None:
                tmp_dir.mkdir(parents=True, exist_ok=True)
                digest = hashlib.sha256(str(source).encode("utf-8")).hexdigest()[:8]
                output = tmp_dir / f"{source.stem}_{digest}.png"
                try:
                    with Image.open(source) as image:
                        image.seek(0)
                        normalized = ImageOps.exif_transpose(image)
                        has_alpha = (
                            normalized.mode in {"RGBA", "LA"}
                            or "transparency" in normalized.info
                        )
                        normalized = normalized.convert("RGBA" if has_alpha else "RGB")
                        normalized.save(output, "PNG", optimize=True)
                except OSError as exc:
                    log.warning(f"图片转换失败，保留原图: {source} - {exc}")
                    return match.group(0)
                converted[source] = output

            title = match.group("title") or ""
            return f"![{match.group('alt')}](<{output.resolve().as_posix()}>{title})"

        new_text = cls._map_outside_fenced_code(
            text,
            lambda segment: cls._MARKDOWN_IMAGE_RE.sub(replace_image, segment),
        )
        if new_text != text:
            input_path.write_text(new_text, encoding="utf-8")
            log.info(f"已将 {len(converted)} 张图片转换为 PNG: {input_path.name}")
        return [tmp_dir] if converted else []

    # ── Mermaid 图表预处理 ────────────────────────────────

    async def _preprocess_mermaid(
        self,
        input_path: Path,
        *,
        mark_for_docx: bool = False,
    ) -> list[Path]:
        """
        将 Markdown 中的 ```mermaid 代码块渲染为图片

        流程：
        1. 扫描文件中的 ```mermaid ... ``` 代码块
        2. 使用 Edge headless + 内嵌 mermaid.js 渲染为 PNG
        3. 将原代码块替换为 ![](<图片路径>)
        4. 写回文件

        Returns:
            创建的临时目录列表（供后续清理使用）

        """
        created_dirs: list[Path] = []

        try:
            text = input_path.read_text(encoding="utf-8")
        except OSError:
            return created_dirs

        # 匹配 ```mermaid ... ``` 代码块
        pattern = re.compile(
            r"```mermaid\w*[ \t]*\n(.*?)```",
            re.DOTALL,
        )

        matches = list(pattern.finditer(text))
        if not matches:
            return created_dirs

        # 临时目录
        stem_hash = hashlib.sha256(input_path.stem.encode("utf-8")).hexdigest()[:8]
        tmp_dir = input_path.parent / f"_mermaid_{stem_hash}"
        tmp_dir.mkdir(parents=True, exist_ok=True)
        created_dirs.append(tmp_dir)

        # ── 收集所有图表 ──
        diagram_tuples: list[tuple[str, Path]] = []
        for idx, m in enumerate(matches):
            code_content = m.group(1).strip()
            if not code_content:
                continue
            png_file = tmp_dir / f"diagram_{idx}.png"
            diagram_tuples.append((code_content, png_file))

        if not diagram_tuples:
            return created_dirs

        # ── 渲染 ──
        any_rendered = False
        new_text = text

        # 1) 用 Edge headless 渲染 Mermaid 图表
        if mermaid_renderer_available():
            log.info(f"使用 Edge 渲染器处理 {len(diagram_tuples)} 个 Mermaid 图表")
            results = await render_diagrams(diagram_tuples)

            # 倒序替换
            for idx in reversed(range(len(matches))):
                if idx >= len(results):
                    continue
                m = matches[idx]
                code_content = m.group(1).strip()
                if not code_content:
                    start, end = m.start(), m.end()
                    new_text = new_text[:start] + "\n" + new_text[end:]
                    continue
                if results[idx]:
                    png_file = tmp_dir / f"diagram_{idx}.png"
                    abs_path_str = png_file.resolve().as_posix()
                    title = f' "{MERMAID_IMAGE_MARKER}"' if mark_for_docx else ""
                    img_md = f"![](<{abs_path_str}>{title})\n"
                    start, end = m.start(), m.end()
                    new_text = new_text[:start] + img_md + new_text[end:]
                    any_rendered = True
                else:
                    # 渲染失败的保留原文
                    log.warning(f"Mermaid 图表 #{idx} 渲染失败，保留原文")
        else:
            log.warning("Edge 渲染器不可用，跳过 Mermaid 图表渲染")
            return created_dirs

        # ── 写回文件 ──
        if any_rendered and new_text != text:
            try:
                input_path.write_text(new_text, encoding="utf-8")
                log.info(f"Mermaid 预处理完成：{len(matches)} 个图表")
                # 后验证
                mermaid_refs = re.findall(r"!\[\]\(<([^>]+)>", new_text)
                for ref in mermaid_refs:
                    ref_path = Path(ref)
                    if ref_path.exists():
                        log.debug(f"  图片已确认: {ref}")
                    else:
                        log.warning(f"  图片路径无效: {ref}")
            except OSError:
                pass
        elif not any_rendered:
            log.warning("所有 Mermaid 图表渲染均失败，跳过文件修改")

        return created_dirs

    @staticmethod
    def _cleanup_mermaid_dir(tmp_dir: Path) -> None:
        """清理 Mermaid 临时目录"""
        if not tmp_dir.exists():
            return
        try:
            shutil.rmtree(tmp_dir)
            log.debug(f"Mermaid 临时目录已清理: {tmp_dir}")
        except OSError as e:
            log.warning(f"清理 Mermaid 临时目录失败: {tmp_dir} - {e}")

    # ── 表格样式后处理 ─────────────────────────────────────

    def _apply_table_styles(self, docx_path: Path, slug: str) -> None:
        """对 docx 中所有表格应用模板的 table 样式配置"""
        tc = self._template_mgr.get_table_config(slug)
        if not tc:
            log.debug(f"模板 '{slug}' 无表格样式配置，跳过")
            return

        try:
            doc = Document(str(docx_path))
        except Exception as e:
            log.warning(f"无法打开 docx 应用表格样式: {e}")
            return

        tables = doc.tables
        if not tables:
            return

        log.info(f"模板 '{slug}' 表格样式: 共 {len(tables)} 个表格")

        # 字体/字号解析
        font_name = str(tc.get("font", "")).strip() or None
        font_size_pt = _parse_size(tc.get("size", "")) if tc.get("size") else None

        # 段落格式
        align = ALIGN_MAP.get(str(tc.get("alignment", "")).strip())

        # 表头样式
        hdr_font = str(tc.get("header_font", "")).strip() or font_name
        header_size = tc.get("header_size")
        hdr_size_pt = _parse_size(header_size) if header_size else font_size_pt
        hdr_bold = tc.get("header_bold", False)
        hdr_align = ALIGN_MAP.get(str(tc.get("header_alignment", "")).strip(), align)
        hdr_bg = str(tc.get("header_background", "")).strip() or None

        # 表体样式
        body_font = str(tc.get("body_font", "")).strip() or font_name
        body_size_pt = _parse_size(tc.get("body_size", "")) if tc.get("body_size") else font_size_pt
        body_align = ALIGN_MAP.get(str(tc.get("body_alignment", "")).strip(), align)

        # 单元格内边距
        pad = tc.get("cell_padding", {})
        pad_top = pad.get("top", 2)
        pad_bot = pad.get("bottom", 2)
        pad_left = pad.get("left", 4)
        pad_right = pad.get("right", 4)

        # 斑马纹
        stripe = tc.get("stripe_rows", False)
        stripe_color = str(tc.get("stripe_color", "")).strip() or None

        for table in tables:
            rows = table.rows
            if not rows:
                continue

            # 边框样式
            border_style = tc.get("border_style")
            if border_style == "three_line_table":
                _apply_three_line_border(table, tc)
            elif border_style == "grid":
                _apply_grid_border(table, tc)

            # 表格撑满到页面左右边距
            if tc.get("full_width"):
                tbl = table._tbl  # type: ignore[attr-defined]
                tblPr = tbl.find(qn("w:tblPr"))
                if tblPr is None:
                    tblPr = OxmlElement("w:tblPr")
                    tbl.insert(0, tblPr)
                # 移除旧的宽度定义
                for old_w in tblPr.findall(qn("w:tblW")):
                    tblPr.remove(old_w)
                # 设置表格宽度为 100%（5000 = 100% 在 pct 单位下）
                tblW = OxmlElement("w:tblW")
                tblW.set(qn("w:w"), "5000")
                tblW.set(qn("w:type"), "pct")
                tblPr.append(tblW)

            for row_idx, row in enumerate(rows):
                is_header = row_idx == 0
                cells = row.cells

                for cell in cells:
                    # 内边距
                    _set_cell_margins(cell, pad_top, pad_bot, pad_left, pad_right)

                    if is_header:
                        # 表头
                        _apply_cell_text_format(
                            cell,
                            font_name=hdr_font,
                            font_size_pt=hdr_size_pt,
                            bold=hdr_bold,
                            alignment=hdr_align,
                        )
                        if hdr_bg:
                            _set_cell_shading(cell, hdr_bg)
                    else:
                        # 表体
                        _apply_cell_text_format(
                            cell,
                            font_name=body_font,
                            font_size_pt=body_size_pt,
                            bold=None,
                            alignment=body_align,
                        )
                        # 斑马纹
                        if stripe and stripe_color and row_idx % 2 == 0:
                            _set_cell_shading(cell, stripe_color)

        # 在每个表格后插入一个空行（避免表格紧贴下一内容）
        for table in reversed(tables):
            tbl = table._tbl
            # 构建空段落
            new_p = OxmlElement("w:p")
            pPr = OxmlElement("w:pPr")
            # 设置一个合理的行间距（约 6pt 空行）
            spacing = OxmlElement("w:spacing")
            spacing.set(qn("w:line"), "360")
            spacing.set(qn("w:lineRule"), "auto")
            pPr.append(spacing)
            new_p.append(pPr)
            tbl.addnext(new_p)

        try:
            doc.save(str(docx_path))
        except Exception as e:
            log.warning(f"保存 docx 表格样式失败: {e}")
