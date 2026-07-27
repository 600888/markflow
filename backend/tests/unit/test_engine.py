"""Pandoc 引擎单元测试"""

from __future__ import annotations

from pathlib import Path
from unittest.mock import AsyncMock, patch

import pytest
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches
from PIL import Image

from app.core.engine import MERMAID_IMAGE_MARKER, PandocEngine
from app.models import OutputFormat
from app.utils.exceptions import ConversionError


class TestPandocEngine:
    def test_format_docx_mermaid_images_centers_and_scales_to_page(
        self,
        tmp_path: Path,
    ) -> None:
        image_path = tmp_path / "wide.png"
        Image.new("RGB", (1600, 400), "white").save(image_path)

        document = Document()
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.left_indent = Inches(0.4)
        shape = paragraph.add_run().add_picture(str(image_path), width=Inches(10))
        shape._inline.docPr.set("title", MERMAID_IMAGE_MARKER)  # noqa: SLF001
        docx_path = tmp_path / "wide-mermaid.docx"
        document.save(docx_path)

        PandocEngine._format_docx_mermaid_images(docx_path)  # noqa: SLF001

        converted = Document(docx_path)
        section = converted.sections[0]
        available_width = (
            section.page_width - section.left_margin - section.right_margin
        )
        converted_shape = converted.inline_shapes[0]
        assert converted_shape.width == available_width
        assert converted_shape.height == round(available_width / 4)
        assert converted.paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.CENTER
        assert converted.paragraphs[0]._p.pPr.find(qn("w:ind")) is None  # noqa: SLF001

    @patch("pypandoc.get_pandoc_path", return_value="/usr/bin/pandoc")
    async def test_preprocess_mermaid_marks_docx_images(
        self,
        mock_pandoc_path,
        tmp_path: Path,
    ) -> None:
        input_file = tmp_path / "diagram.md"
        input_file.write_text(
            "```mermaid\nflowchart LR\nA --> B\n```\n",
            encoding="utf-8",
        )
        engine = PandocEngine()

        with (
            patch("app.core.engine.mermaid_renderer_available", return_value=True),
            patch(
                "app.core.engine.render_diagrams",
                new=AsyncMock(return_value=[True]),
            ),
        ):
            await engine._preprocess_mermaid(  # noqa: SLF001
                input_file,
                mark_for_docx=True,
            )

        markdown = input_file.read_text(encoding="utf-8")
        assert f'"{MERMAID_IMAGE_MARKER}"' in markdown
        assert markdown.startswith("![](<")

    def test_get_toc_depth(self) -> None:
        assert PandocEngine._get_toc_depth(["--toc", "--toc-depth", "4"]) == 4  # noqa: SLF001
        assert PandocEngine._get_toc_depth(["--toc-depth=2"]) == 2  # noqa: SLF001
        assert PandocEngine._get_toc_depth(["--toc-depth", "invalid"]) == 3  # noqa: SLF001

    def test_internal_page_metadata_is_read_and_removed(self) -> None:
        args = [
            "--metadata",
            "title=报告",
            "--metadata",
            "markflow-title-page=true",
            "--metadata=markflow-page-header=季度报告",
        ]
        value = PandocEngine._get_metadata_value(  # noqa: SLF001
            args,
            "markflow-title-page",
        )
        assert value == "true"
        filtered = PandocEngine._remove_metadata_keys(  # noqa: SLF001
            args,
            {"markflow-title-page", "markflow-page-header"},
        )
        assert filtered == ["--metadata", "title=报告"]

    def test_apply_docx_title_page_and_header(self, tmp_path: Path) -> None:
        document = Document()
        document.add_paragraph("季度总结", style="Title")
        document.add_paragraph("张三", style="Subtitle")
        document.add_heading("第一章", level=1)
        output = tmp_path / "page-options.docx"
        document.save(output)

        PandocEngine._apply_docx_page_options(  # noqa: SLF001
            output,
            title_page=True,
            page_header="季度报告",
        )

        converted = Document(output)
        assert converted.sections[0].different_first_page_header_footer is False
        assert converted.sections[0].header.paragraphs[0].text == "季度报告"
        assert (
            converted.sections[0].header.paragraphs[0].alignment
            == WD_ALIGN_PARAGRAPH.CENTER
        )
        header_run = converted.sections[0].header.paragraphs[0].runs[0]
        assert header_run.font.name == "宋体"
        assert header_run.font.size.pt == 10.5
        run_fonts = header_run._r.rPr.find(qn("w:rFonts"))  # noqa: SLF001
        assert run_fonts.get(qn("w:eastAsia")) == "宋体"
        bottom_border = converted.sections[0].header.paragraphs[0]._p.find(  # noqa: SLF001
            "./" + qn("w:pPr") + "/" + qn("w:pBdr") + "/" + qn("w:bottom"),
        )
        assert bottom_border is not None
        assert bottom_border.get(qn("w:val")) == "single"
        assert bottom_border.get(qn("w:color")) == "000000"
        assert bottom_border.get(qn("w:space")) == "1"
        assert converted.paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.CENTER
        assert converted.paragraphs[1].alignment == WD_ALIGN_PARAGRAPH.CENTER
        assert not converted.paragraphs[1]._p.findall(  # noqa: SLF001
            ".//" + qn("w:br")
        )
        title_page_break = converted.paragraphs[2]._p.find(  # noqa: SLF001
            "./" + qn("w:pPr") + "/" + qn("w:pageBreakBefore"),
        )
        assert title_page_break is not None
        assert title_page_break.get(qn("w:val")) == "1"

    def test_apply_docx_header_uses_template_config(self, tmp_path: Path) -> None:
        document = Document()
        output = tmp_path / "custom-header.docx"
        document.save(output)

        PandocEngine._apply_docx_page_options(  # noqa: SLF001
            output,
            title_page=False,
            page_header="自定义页眉",
            header_config={
                "font": "微软雅黑",
                "size": "小四",
                "alignment": "right",
                "border_bottom": {"weight": 1.5, "color": "#123456"},
            },
        )

        converted = Document(output)
        paragraph = converted.sections[0].header.paragraphs[0]
        assert paragraph.alignment == WD_ALIGN_PARAGRAPH.RIGHT
        assert paragraph.runs[0].font.name == "微软雅黑"
        assert paragraph.runs[0].font.size.pt == 12
        bottom_border = paragraph._p.find(  # noqa: SLF001
            "./" + qn("w:pPr") + "/" + qn("w:pBdr") + "/" + qn("w:bottom"),
        )
        assert bottom_border is not None
        assert bottom_border.get(qn("w:sz")) == "12"
        assert bottom_border.get(qn("w:color")) == "123456"

    def test_populate_docx_toc_cache(self, tmp_path: Path) -> None:
        document = Document()
        body = document.element.body

        sdt = OxmlElement("w:sdt")
        sdt_properties = OxmlElement("w:sdtPr")
        doc_part = OxmlElement("w:docPartObj")
        gallery = OxmlElement("w:docPartGallery")
        gallery.set(qn("w:val"), "Table of Contents")
        doc_part.append(gallery)
        sdt_properties.append(doc_part)
        sdt.append(sdt_properties)

        content = OxmlElement("w:sdtContent")
        title = OxmlElement("w:p")
        title_run = OxmlElement("w:r")
        title_text = OxmlElement("w:t")
        title_text.text = "目录"
        title_run.append(title_text)
        title.append(title_run)
        content.append(title)

        field = OxmlElement("w:p")
        field_run = OxmlElement("w:r")
        begin = OxmlElement("w:fldChar")
        begin.set(qn("w:fldCharType"), "begin")
        field_run.append(begin)
        instruction = OxmlElement("w:instrText")
        instruction.text = 'TOC \\o "1-2" \\h \\z \\u'
        field_run.append(instruction)
        separate = OxmlElement("w:fldChar")
        separate.set(qn("w:fldCharType"), "separate")
        field_run.append(separate)
        end = OxmlElement("w:fldChar")
        end.set(qn("w:fldCharType"), "end")
        field_run.append(end)
        field.append(field_run)
        content.append(field)
        sdt.append(content)
        body.insert(0, sdt)

        heading1 = document.add_heading("第一章", level=1)._p  # noqa: SLF001
        bookmark1 = OxmlElement("w:bookmarkStart")
        bookmark1.set(qn("w:id"), "1")
        bookmark1.set(qn("w:name"), "chapter-one")
        body.insert(body.index(heading1), bookmark1)
        document.add_heading("第一节", level=2)
        document.add_heading("不应进入目录", level=3)

        output = tmp_path / "toc.docx"
        document.save(output)

        PandocEngine._populate_docx_toc_cache(output, 2)  # noqa: SLF001

        converted = Document(output)
        converted_content = converted.element.body.find(
            qn("w:sdt") + "/" + qn("w:sdtContent"),
        )
        assert converted_content is not None
        visible_text = [
            "".join(node.text or "" for node in paragraph.iter(qn("w:t")))
            for paragraph in converted_content.findall(qn("w:p"))
        ]
        assert "第一章" in visible_text
        assert "第一节" in visible_text
        assert "不应进入目录" not in visible_text
        update_fields = converted.settings.element.find(qn("w:updateFields"))
        assert update_fields is not None
        assert update_fields.get(qn("w:val")) == "true"

        toc_container = converted_content.getparent()
        assert toc_container is not None
        body = converted.element.body
        paragraphs_after_toc = [
            element
            for element in list(body)[body.index(toc_container) + 1 :]
            if element.tag == qn("w:p")
        ]
        assert paragraphs_after_toc
        page_break = paragraphs_after_toc[0].find(
            "./" + qn("w:pPr") + "/" + qn("w:pageBreakBefore"),
        )
        assert page_break is not None
        assert page_break.get(qn("w:val")) == "1"
        assert paragraphs_after_toc[0].find(".//" + qn("w:br")) is None

    @patch("pypandoc.get_pandoc_path", return_value="/usr/bin/pandoc")
    async def test_convert_populates_visible_toc(
        self,
        mock_pandoc_path,
        tmp_path: Path,
    ) -> None:
        input_file = tmp_path / "test.md"
        input_file.write_text("# 第一章", encoding="utf-8")
        output_file = input_file.with_suffix(".docx")
        output_file.write_text("mock output", encoding="utf-8")

        engine = PandocEngine()
        with (
            patch("pypandoc.convert_file", return_value=str(output_file)),
            patch.object(engine, "_populate_docx_toc_cache") as populate_toc,
        ):
            await engine.convert(
                input_file,
                OutputFormat.DOCX,
                extra_args=["--toc", "--toc-depth", "2"],
            )

        populate_toc.assert_called_once_with(output_file, 2)

    @patch("pypandoc.get_pandoc_path", return_value="/usr/bin/pandoc")
    async def test_convert_can_skip_mermaid_preprocessing(
        self,
        mock_pandoc_path,
        tmp_path: Path,
    ) -> None:
        input_file = tmp_path / "test.md"
        input_file.write_text("```mermaid\ngraph TD\nA-->B\n```", encoding="utf-8")
        output_file = input_file.with_suffix(".html")
        output_file.write_text("mock output", encoding="utf-8")

        engine = PandocEngine()
        with (
            patch("pypandoc.convert_file", return_value=str(output_file)),
            patch.object(engine, "_preprocess_mermaid") as preprocess_mermaid,
        ):
            await engine.convert(
                input_file,
                OutputFormat.HTML,
                convert_mermaid=False,
            )

        preprocess_mermaid.assert_not_called()

    def test_strip_images_when_disabled(self, tmp_path: Path) -> None:
        input_file = tmp_path / "test.md"
        input_file.write_text(
            "正文 ![示意图](image.png)\n"
            '<img src="other.png" alt="另一张图">\n'
            "```markdown\n![代码示例](keep.png)\n```\n",
            encoding="utf-8",
        )

        created_dirs = PandocEngine._preprocess_images(  # noqa: SLF001
            input_file,
            convert_images=False,
        )

        content = input_file.read_text(encoding="utf-8")
        assert created_dirs == []
        assert "正文 示意图" in content
        assert "另一张图" in content
        assert "other.png" not in content
        assert "![代码示例](keep.png)" in content

    def test_convert_local_bmp_to_png(self, tmp_path: Path) -> None:
        source = tmp_path / "sample.bmp"
        Image.new("RGB", (4, 3), color=(20, 80, 140)).save(source)
        input_file = tmp_path / "test.md"
        input_file.write_text(
            '![示意图](sample.bmp "标题")',
            encoding="utf-8",
        )

        created_dirs = PandocEngine._preprocess_images(  # noqa: SLF001
            input_file,
            convert_images=True,
        )

        assert len(created_dirs) == 1
        png_files = list(created_dirs[0].glob("*.png"))
        assert len(png_files) == 1
        assert png_files[0].is_file()
        assert ".png>" in input_file.read_text(encoding="utf-8")
        with Image.open(png_files[0]) as converted:
            assert converted.format == "PNG"

    @patch("pypandoc.get_pandoc_path", return_value="/usr/bin/pandoc")
    async def test_all_formats_supported(self, mock_pandoc) -> None:
        """所有 OutputFormat 枚举值都应该被支持"""
        engine = PandocEngine()
        for fmt in OutputFormat:
            assert await engine.validate_format(fmt), f"{fmt} 应该被支持"

    @patch("pypandoc.get_pandoc_path", return_value="/usr/bin/pandoc")
    async def test_convert_success(self, mock_pandoc_path, tmp_path: Path) -> None:
        """正常转换应返回 ConversionResult"""
        input_file = tmp_path / "test.md"
        input_file.write_text("# Hello")
        output_file = input_file.with_suffix(".docx")
        output_file.write_text("mock output")

        engine = PandocEngine()
        with patch("pypandoc.convert_file", return_value=str(output_file)):
            result = await engine.convert(input_file, OutputFormat.DOCX)

        assert result.output_format == OutputFormat.DOCX
        assert result.duration_ms >= 0
        assert result.file_size >= 0

    @patch("pypandoc.get_pandoc_path", return_value="/usr/bin/pandoc")
    @patch("pypandoc.convert_file", side_effect=RuntimeError("Pandoc crashed"))
    async def test_convert_failure(self, mock_convert, mock_pandoc_path, tmp_path: Path) -> None:
        """Pandoc 异常应转为 ConversionError"""
        input_file = tmp_path / "test.md"
        input_file.write_text("# Hello")

        engine = PandocEngine()
        with pytest.raises(ConversionError) as exc_info:
            await engine.convert(input_file, OutputFormat.DOCX)

        assert "Pandoc" in exc_info.value.message

    @patch("pypandoc.get_pandoc_path", return_value="/usr/bin/pandoc")
    async def test_convert_with_progress(self, mock_pandoc_path, tmp_path: Path) -> None:
        """应调用进度回调"""
        input_file = tmp_path / "test.md"
        input_file.write_text("# Hello")
        output_file = input_file.with_suffix(".docx")
        output_file.write_text("mock output")

        on_progress = AsyncMock()

        engine = PandocEngine()
        with patch("pypandoc.convert_file", return_value=str(output_file)):
            await engine.convert(input_file, OutputFormat.DOCX, on_progress=on_progress)

        assert on_progress.await_count >= 1
