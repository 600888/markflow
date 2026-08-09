"""Tests for Word reference template generation."""

from io import BytesIO
from pathlib import Path

from docx import Document

from app.services.template_generator import TemplateGenerator


def test_heading4_is_not_italic_when_omitted() -> None:
    generator = TemplateGenerator()

    result = generator.generate_reference(
        {
            "body": {
                "font": "Arial",
                "size": 11,
            },
        },
    )

    document = Document(BytesIO(result))
    assert document.styles["Heading 4"].font.italic is False


def test_heading4_can_be_explicitly_italic() -> None:
    generator = TemplateGenerator()

    result = generator.generate_reference(
        {
            "heading4": {
                "font": "Arial",
                "size": 11,
                "italic": True,
            },
        },
    )

    document = Document(BytesIO(result))
    assert document.styles["Heading 4"].font.italic is True


def test_heading5_style_is_generated() -> None:
    generator = TemplateGenerator()

    result = generator.generate_reference(
        {
            "heading5": {
                "font": "Arial",
                "size": 10.5,
                "bold": True,
                "italic": False,
            },
        },
    )

    document = Document(BytesIO(result))
    heading5 = document.styles["Heading 5"]
    assert heading5.font.name == "Arial"
    assert heading5.font.bold is True
    assert heading5.font.italic is False


def test_custom_template_can_be_created_updated_and_deleted(tmp_path: Path) -> None:
    generator = TemplateGenerator(tmp_path)

    generator.save_custom_template(
        name="First name",
        slug="team-report",
        styles_config={"body": {"font": "Arial", "size": 11}},
    )
    assert generator.custom_template_exists("team-report")

    generator.save_custom_template(
        name="Updated name",
        slug="team-report",
        styles_config={"body": {"font": "Calibri", "size": 12}},
    )
    templates = generator.list_custom_templates()
    assert templates[0]["name"] == "Updated name"

    assert generator.delete_custom_template("team-report") is True
    assert generator.custom_template_exists("team-report") is False


def test_custom_template_rejects_unsafe_slug(tmp_path: Path) -> None:
    generator = TemplateGenerator(tmp_path)

    assert generator.custom_template_exists("../builtin") is False
    assert generator.delete_custom_template("../builtin") is False
